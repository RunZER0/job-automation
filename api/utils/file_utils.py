import io
import os
import tempfile
from fastapi import UploadFile, HTTPException
from typing import Union

async def extract_text_from_file(file: UploadFile) -> str:
    """Extract text from uploaded file (supports .txt, .pdf, and .docx)"""
    
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")
    
    filename = file.filename.lower()
    
    # Read file content once
    try:
        content = await file.read()
        # Reset file pointer for potential re-reading
        await file.seek(0)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading file: {str(e)}")

    # Handle PDF files
    if filename.endswith('.pdf'):
        try:
            from pdfminer.high_level import extract_text
            
            # Save the uploaded PDF to a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            
            text = extract_text(tmp_path)
            
            # Clean up temporary file
            try:
                os.unlink(tmp_path)
            except OSError:
                pass  # File might already be deleted
            
            if not text.strip():
                raise HTTPException(status_code=400, detail="No text found in PDF file")
            
            return text.strip()
            
        except ImportError:
            raise HTTPException(status_code=500, detail="pdfminer.six library not installed. Please install it with: pip install pdfminer.six")
        except Exception as e:
            # Clean up temporary file if it exists
            if 'tmp_path' in locals():
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass
            raise HTTPException(status_code=400, detail=f"PDF extraction failed: {str(e)}")

    # Handle DOCX files
    elif filename.endswith('.docx'):
        try:
            from docx import Document
            
            # Save the uploaded DOCX to a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            
            doc = Document(tmp_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            # Clean up temporary file
            try:
                os.unlink(tmp_path)
            except OSError:
                pass  # File might already be deleted
            
            if not text.strip():
                raise HTTPException(status_code=400, detail="No text found in DOCX file")
            
            return text.strip()
            
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx library not installed. Please install it with: pip install python-docx")
        except Exception as e:
            # Clean up temporary file if it exists
            if 'tmp_path' in locals():
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass
            raise HTTPException(status_code=400, detail=f"DOCX extraction failed: {str(e)}")

    # Handle TXT files
    elif filename.endswith('.txt'):
        try:
            # Try UTF-8 first
            text = content.decode('utf-8')
            return text
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                text = content.decode('latin-1', errors='ignore')
                return text
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error decoding text file: {str(e)}")

    # Unsupported file types
    else:
        supported_types = ['.pdf', '.docx', '.txt']
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type. Please upload one of: {', '.join(supported_types)}"
        )

# Synchronous version for backward compatibility
def extract_text_from_file_sync(file: UploadFile) -> str:
    """Synchronous version of extract_text_from_file (deprecated - use async version)"""
    
    if not file.filename:
        return "Error: No filename provided"
    
    filename = file.filename.lower()
    
    # Handle PDF files
    if filename.endswith('.pdf'):
        try:
            from pdfminer.high_level import extract_text
            
            # Read file content
            content = file.file.read()
            
            # Save the uploaded PDF to a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            
            text = extract_text(tmp_path)
            os.unlink(tmp_path)
            return text.strip() if text.strip() else "No text found in PDF"
            
        except Exception as e:
            return f"PDF extraction failed: {str(e)}"

    # Handle DOCX files
    elif filename.endswith('.docx'):
        try:
            from docx import Document
            
            # Read file content
            content = file.file.read()
            
            # Save the uploaded DOCX to a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            
            doc = Document(tmp_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            os.unlink(tmp_path)
            return text.strip() if text.strip() else "No text found in DOCX"
            
        except Exception as e:
            return f"DOCX extraction failed: {str(e)}"

    # Handle TXT files
    elif filename.endswith('.txt'):
        content = file.file.read()
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return content.decode('latin-1', errors='ignore')
            except Exception as e:
                return f"Error decoding text file: {str(e)}"

    # Fallback for unknown file types
    else:
        return "Unsupported file type. Please upload PDF, DOCX, or TXT."

